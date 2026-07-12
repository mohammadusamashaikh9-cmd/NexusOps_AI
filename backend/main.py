from fastapi import FastAPI, Form, File, UploadFile, HTTPException
from pydantic import BaseModel
from typing import Any, Dict

app = FastAPI(
    title="NexusOps AI API",
    description="Backend API for the NexusOps AI operations platform",
    version="0.1.0"
)

from fastapi.middleware.cors import CORSMiddleware
allowed_origins = [
    "http://localhost:5173",
    "http://localhost:3000",
    "https://nexus-ops-ai-umber.vercel.app",
    "https://nexus-ops-ai.vercel.app"
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
def root():
    return {"status": "ok", "service": "NexusOps AI backend"}

@app.get("/api/health")
def health_check():
    return {"status": "ok", "message": "NexusOps AI backend is running."}

# Mock Endpoints for MVP Agents
from services.workflow_service import run_workflow

import io
import pypdf
import docx

def extract_text_from_bytes(filename: str, content_bytes: bytes) -> str:
    filename_lower = filename.lower()
    
    if filename_lower.endswith(".pdf"):
        try:
            reader = pypdf.PdfReader(io.BytesIO(content_bytes))
            text = "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
            return text.strip()
        except Exception as e:
            raise ValueError("Failed to extract text from PDF file.")
            
    elif filename_lower.endswith(".docx"):
        try:
            doc = docx.Document(io.BytesIO(content_bytes))
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            return text.strip()
        except Exception as e:
            raise ValueError("Failed to extract text from DOCX file.")
            
    elif filename_lower.endswith(".txt") or filename_lower.endswith(".md"):
        try:
            return content_bytes.decode('utf-8').strip()
        except UnicodeDecodeError:
            raise ValueError("File must be valid UTF-8 text.")
            
    else:
        raise ValueError("Unsupported file format. Only PDF, DOCX, TXT, and MD are supported.")

@app.post("/api/workflow/run")
async def api_run_workflow(
    task: str = Form(...),
    document: UploadFile = File(None)
):
    document_context = None
    document_name = None
    
    if document is not None and document.filename:
        filename_lower = document.filename.lower()
        if not (filename_lower.endswith(".txt") or filename_lower.endswith(".md") or filename_lower.endswith(".pdf") or filename_lower.endswith(".docx")):
            raise HTTPException(status_code=400, detail="Only PDF, DOCX, TXT, and MD files are supported.")
        
        content_bytes = await document.read(1024 * 1024 + 1)
        if len(content_bytes) > 1024 * 1024:
            raise HTTPException(status_code=400, detail="File size exceeds the 1MB limit.")
            
        try:
            document_context = extract_text_from_bytes(document.filename, content_bytes)
            if not document_context:
                raise ValueError("No extractable text found in the document.")
        except ValueError as ve:
            raise HTTPException(status_code=400, detail=str(ve))
            
        document_name = document.filename
        
    return run_workflow(task, document_context, document_name)

from fastapi.responses import Response

@app.post("/api/report/export")
async def api_export_report(
    report_content: str = Form(...),
    format: str = Form(...)
):
    if not report_content.strip():
        raise HTTPException(status_code=400, detail="Report content cannot be empty.")
    
    if format == 'md':
        return Response(
            content=report_content, 
            media_type="text/markdown", 
            headers={"Content-Disposition": "attachment; filename=nexusops-report.md"}
        )
        
    elif format == 'pdf':
        try:
            from fpdf import FPDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Helvetica", size=11)
            safe_content = report_content.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 6, text=safe_content)
            pdf_bytes = pdf.output()
            return Response(
                content=bytes(pdf_bytes), 
                media_type="application/pdf", 
                headers={"Content-Disposition": "attachment; filename=nexusops-report.pdf"}
            )
        except Exception as e:
            raise HTTPException(status_code=500, detail="Failed to generate PDF export.")
            
    elif format == 'docx':
        try:
            doc = docx.Document()
            for line in report_content.split('\n'):
                doc.add_paragraph(line)
            buf = io.BytesIO()
            doc.save(buf)
            return Response(
                content=buf.getvalue(), 
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                headers={"Content-Disposition": "attachment; filename=nexusops-report.docx"}
            )
        except Exception as e:
            raise HTTPException(status_code=500, detail="Failed to generate DOCX export.")
            
    else:
        raise HTTPException(status_code=400, detail="Unsupported export format. Supported: md, pdf, docx.")
